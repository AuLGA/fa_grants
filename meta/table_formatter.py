# A simple script to open all docx files in "QUARTO_PROJECT_OUTPUT_FILES" directory and add a page break before each heading.

import os
import docx
from glob import glob
from docx.shared import Pt

files = ["./modelling_outcomes_binned.docx", "./aclg_note.docx"]

for file in files:
    doc = docx.Document(file)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading 1"):
            run = paragraph.insert_paragraph_before().add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(8)

    doc.save(file)

