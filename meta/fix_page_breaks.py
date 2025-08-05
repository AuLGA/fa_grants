# A simple script to open all docx files in "QUARTO_PROJECT_OUTPUT_FILES" directory and add a page break before each heading.

import os
import docx
from glob import glob

output_dir = os.getenv("QUARTO_PROJECT_OUTPUT_FILES")

output_files = output_dir.split("\n")

for file in output_files:
    if not file.endswith(".docx"):
        continue
    print(f"Processing {file}...")

    doc = docx.Document(file)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading 1"):

            print(f"Adding page break before {paragraph.text}...")

            # Insert a page break before the heading

            run = paragraph.insert_paragraph_before().add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)
            
    doc.save(file)
    print(f"Saved {file} with page breaks added.")

